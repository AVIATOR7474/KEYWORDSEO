import streamlit as st
import docx
from docx import Document
import io
import html
import re
from collections import defaultdict, Counter
import spacy
from statistics import mean
import pandas as pd

# محاولة تحميل نموذج اللغة العربية
import subprocess
import importlib.util

def ensure_arabic_model():
    try:
        return spacy.load("ar_core_news_sm")
    except OSError:
        with st.spinner("📥 تحميل نموذج اللغة العربية..."):
            subprocess.run(["python", "-m", "spacy", "download", "ar_core_news_sm"])
            return spacy.load("ar_core_news_sm")

nlp = ensure_arabic_model()


st.set_page_config(page_title="SEO Advanced Document Optimizer", layout="wide", page_icon="🔍")

# CSS مخصص
st.markdown("""
<style>
    .header { color: #2E86AB; font-size: 2.2rem; font-weight: bold; }
    .subheader { color: #A23B72; font-size: 1.4rem; margin-top: 1.5rem; }
    .metric-box { background-color: #F8F9FA; border-radius: 10px; padding: 15px; margin-bottom: 15px; }
    .progress-bar { height: 10px; border-radius: 5px; background-color: #E9ECEF; }
    .progress-fill { height: 100%; border-radius: 5px; background-color: #4CAF50; }
    .warning { color: #FFC107; font-weight: bold; }
    .error { color: #DC3545; font-weight: bold; }
</style>
""", unsafe_allow_html=True)

def main():
    st.markdown('<p class="header">🔍 SEO Advanced Document Optimizer 2025</p>', unsafe_allow_html=True)

    with st.sidebar:
        st.image("https://via.placeholder.com/150x50?text=SEO+2025", width=150)
        st.markdown("### إعدادات التحليل المتقدم")
        analysis_depth = st.selectbox("عمق التحليل", ["سريع", "متوسط", "مفصل"], index=1)
        check_ux = st.checkbox("فحص تجربة المستخدم", True)
        check_semantic = st.checkbox("تحليل السياق الدلالي", True)
        generate_meta = st.checkbox("إنشاء ميتاداتا تلقائية", True)

    tab1, tab2, tab3 = st.tabs(["📤 تحميل الملف", "⚙️ الإعدادات", "📊 النتائج"])

    with tab1:
        uploaded_file = st.file_uploader("رفع ملف DOCX", type=["docx"])
        if uploaded_file:
            with st.expander("معاينة الملف الأولية"):
                doc = Document(uploaded_file)
                st.text("\n".join([p.text for p in doc.paragraphs[:5]]))

    with tab2:
        if uploaded_file:
            col1, col2 = st.columns(2)
            with col1:
                st.markdown('<p class="subheader">الكلمات المفتاحية الرئيسية</p>', unsafe_allow_html=True)
                primary_keywords = st.text_area("(كل سطر يحتوي على كلمة/عبارة)", height=150)
            with col2:
                st.markdown('<p class="subheader">الكلمات المفتاحية الثانوية</p>', unsafe_allow_html=True)
                secondary_keywords = st.text_area("(كل سطر يحتوي على كلمة/عبارة)", height=150)

            st.markdown('<p class="subheader">إعدادات الهيكلة</p>', unsafe_allow_html=True)
            h1_max = st.slider("الحد الأقصى لعناوين H1", 1, 3, 1)
            h2_max = st.slider("الحد الأقصى لعناوين H2 لكل كلمة", 2, 5, 3)
            h3_max = st.slider("الحد الأقصى لعناوين H3 لكل كلمة", 1, 4, 2)

    with tab3:
        if uploaded_file and st.button("بدء التحليل والتحسين", type="primary"):
            with st.spinner("جاري التحليل المتعمق..."):
                results = process_document(uploaded_file, primary_keywords, secondary_keywords, 
                                           h1_max, h2_max, h3_max, analysis_depth, check_ux, check_semantic, generate_meta)
                display_results(results)

def process_document(uploaded_file, primary_keywords, secondary_keywords, h1_max, h2_max, h3_max, 
                     analysis_depth, check_ux, check_semantic, generate_meta):
    doc = Document(uploaded_file)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    primary_kw_list = [kw.strip() for kw in primary_keywords.split("\n") if kw.strip()]
    secondary_kw_list = [kw.strip() for kw in secondary_keywords.split("\n") if kw.strip()]

    results = {
        "original_doc": doc,
        "full_text": full_text,
        "primary_keywords": primary_kw_list,
        "secondary_keywords": secondary_kw_list
    }

    optimized_doc = optimize_headings(doc, primary_kw_list, secondary_kw_list, h1_max, h2_max, h3_max)
    results["optimized_doc"] = optimized_doc

    if check_ux:
        results["ux_analysis"] = analyze_content_health(full_text)

    if check_semantic and nlp:
        results["semantic_analysis"] = analyze_semantic_context(full_text, primary_kw_list + secondary_kw_list)

    if generate_meta:
        results["metadata"] = generate_metadata(full_text, primary_kw_list)

    return results

def optimize_headings(doc, primary_keywords, secondary_keywords, h1_max, h2_max, h3_max):
    optimized_doc = Document()
    kw_usage = defaultdict(lambda: {"h1": 0, "h2": 0, "h3": 0})

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        used_kw = None
        heading_level = None

        for kw in primary_keywords:
            if kw in text:
                if kw_usage[kw]["h1"] < h1_max:
                    heading_level = 1
                elif kw_usage[kw]["h2"] < h2_max:
                    heading_level = 2
                elif kw_usage[kw]["h3"] < h3_max:
                    heading_level = 3
                if heading_level:
                    used_kw = kw
                    kw_usage[kw][f"h{heading_level}"] += 1
                    break

        if not used_kw:
            for kw in secondary_keywords:
                if kw in text:
                    if kw_usage[kw]["h2"] < h2_max:
                        heading_level = 2
                    elif kw_usage[kw]["h3"] < h3_max:
                        heading_level = 3
                    if heading_level:
                        used_kw = kw
                        kw_usage[kw][f"h{heading_level}"] += 1
                        break

        if used_kw and heading_level:
            optimized_doc.add_heading(used_kw, level=heading_level)
            if text != used_kw:
                optimized_doc.add_paragraph(text)
        else:
            optimized_doc.add_paragraph(text)

    return optimized_doc

def analyze_content_health(text):
    sentences = [s.strip() for s in re.split(r'[.!?]', text) if s.strip()]
    words = [word for s in sentences for word in s.split()]
    avg_sentence_len = mean(len(s.split()) for s in sentences) if sentences else 0
    long_sentences = sum(1 for s in sentences if len(s.split()) > 25)

    return {
        "word_count": len(words),
        "sentence_count": len(sentences),
        "avg_sentence_length": round(avg_sentence_len, 1),
        "long_sentences": long_sentences,
        "readability": "جيدة" if 15 <= avg_sentence_len <= 25 else "تحتاج تحسين"
    }

def analyze_semantic_context(text, keywords):
    if not nlp:
        return {"error": "نموذج اللغة غير محمل"}

    doc = nlp(text)
    semantic_relations = defaultdict(list)

    for kw in keywords:
        if kw in text:
            kw_doc = nlp(kw)
            for sent in doc.sents:
                if kw_doc.text in sent.text:
                    semantic_relations[kw].append([
                        token.text for token in sent 
                        if token.is_alpha and token.text.lower() not in [k.lower() for k in keywords]
                    ][:5])

    return {
        "top_related_words": {
            kw: Counter([word for lst in semantic_relations[kw] for word in lst]).most_common(3)
            for kw in semantic_relations
        }
    }

def generate_metadata(text, primary_keywords):
    summary = text[:300].replace("\n", " ") + "..." if len(text) > 300 else text
    main_keyword = primary_keywords[0] if primary_keywords else ""

    return {
        "meta_title": f"{main_keyword} | دليل شامل" if main_keyword else "مقال متكامل",
        "meta_description": summary,
        "focus_keyword": main_keyword,
        "tags": primary_keywords[:5]
    }

def display_results(results):
    st.markdown('<p class="header">📊 نتائج التحليل المتقدم</p>', unsafe_allow_html=True)

    with st.expander("🔍 تحليل الهيكل والعناوين", expanded=True):
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("### الكلمات المفتاحية المستخدمة")
            kw_df = pd.DataFrame({
                "الكلمة": results["primary_keywords"] + results["secondary_keywords"],
                "النوع": ["رئيسية"]*len(results["primary_keywords"]) + ["ثانوية"]*len(results["secondary_keywords"])
            })
            st.dataframe(kw_df, hide_index=True)

        with col2:
            st.markdown("### إحصائيات المحتوى")
            st.metric("عدد الكلمات", results["ux_analysis"]["word_count"])
            st.metric("متوسط طول الجملة", results["ux_analysis"]["avg_sentence_length"])
            st.metric("الجمل الطويلة", results["ux_analysis"]["long_sentences"])

    if "ux_analysis" in results:
        with st.expander("📈 تحليل تجربة المستخدم (UX)"):
            progress_html = f"""
            <div class="metric-box">
                <h4>قابلية القراءة: <span class="{'warning' if results['ux_analysis']['readability'] == 'تحتاج تحسين' else ''}">
                {results['ux_analysis']['readability']}</span></h4>
                <div class="progress-bar">
                    <div class="progress-fill" style="width: {min(100, results['ux_analysis']['avg_sentence_length']*4)}%"></div>
                </div>
                <p>المجال المثالي: 15-25 كلمة لكل جملة</p>
            </div>
            """
            st.markdown(progress_html, unsafe_allow_html=True)

    if "semantic_analysis" in results and "top_related_words" in results["semantic_analysis"]:
        with st.expander("🧠 التحليل الدلالي"):
            for kw, related in results["semantic_analysis"]["top_related_words"].items():
                st.markdown(f"**الكلمة المفتاحية:** `{kw}`")
                st.markdown("**الكلمات المرتبطة:** " + "، ".join([f"{w[0]} ({w[1]})" for w in related]))

    if "metadata" in results:
        with st.expander("🏷️ الميتاداتا المولدة"):
            st.text_input("عنوان الصفحة (Meta Title)", results["metadata"]["meta_title"])
            st.text_area("وصف الصفحة (Meta Description)", results["metadata"]["meta_description"])
            st.text_input("الكلمة المفتاحية الرئيسية", results["metadata"]["focus_keyword"])
            st.text_input("الوسوم", ", ".join(results["metadata"]["tags"]))

    with st.expander("💾 تحميل الملف المحسن"):
        col1, col2 = st.columns(2)

        with col1:
            buffer_docx = io.BytesIO()
            results["optimized_doc"].save(buffer_docx)
            buffer_docx.seek(0)
            st.download_button(
                "تحميل DOCX",
                buffer_docx.getvalue(),
                "seo_optimized.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            html_content = docx_to_html(results["optimized_doc"])
            st.download_button(
                "تحميل HTML",
                html_content,
                "seo_optimized.html",
                "text/html"
            )

def docx_to_html(doc):
    html_content = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            html_content.append(f"<h{level} dir='rtl'>{html.escape(text)}</h{level}>")
        else:
            html_content.append(f"<p dir='rtl'>{html.escape(text)}</p>")

    return f"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>مستند محسن لـSEO</title>
</head>
<body>
{"".join(html_content)}
</body>
</html>"""

if __name__ == "__main__":
    main()
